"""
Testni Word (.docx) qilib chiqarish.
Formulalar/asl bezak AYNAN nusxalanadi (matnga aylantirilmaydi).
Faqat: savol qayta raqamlanadi, variantlar aralashtirilib A) B) C)... deb qayta belgilanadi.
Oddiy matn Times New Roman bo'ladi; formulalar (Cambria Math) o'z holicha qoladi.
"""
import io
import re
import random
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
V = "urn:schemas-microsoft-com:vml"
LETTERS = [chr(65 + i) for i in range(26)]


def _ln(el):
    return etree.QName(el).localname if isinstance(el.tag, str) else ""


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), "Times New Roman")


def _force_tnr(el):
    """oddiy run'larga Times New Roman; formulaga (oMath) tegmaymiz."""
    t = _ln(el)
    if t in ("oMath", "oMathPara"):
        return
    if t == "r":
        rpr = el.find("{%s}rPr" % W)
        if rpr is None:
            rpr = etree.SubElement(el, "{%s}rPr" % W)
            el.insert(0, rpr)
        rf = rpr.find("{%s}rFonts" % W)
        if rf is None:
            rf = etree.SubElement(rpr, "{%s}rFonts" % W)
            rpr.insert(0, rf)
        for a in ("ascii", "hAnsi", "cs", "eastAsia"):
            rf.set("{%s}%s" % (W, a), "Times New Roman")
    for c in el:
        _force_tnr(c)


def _append_fragments(doc, p, xml_list, get_media):
    """Saqlangan XML bo'laklarni paragrafga tiklaydi (formulalar aynan)."""
    for s in xml_list:
        # DBMEDIA rasm tokeni bormi?
        mids = re.findall(r"DBMEDIA:(\d+)", s)
        if mids:
            for mid in mids:
                data = get_media(int(mid))
                if data:
                    try:
                        p.add_run().add_picture(io.BytesIO(data), width=Inches(2.2))
                    except Exception:
                        pass
            continue
        try:
            el = etree.fromstring(s)
        except Exception:
            continue
        _force_tnr(el)
        p._p.append(el)


def build(questions, title, get_media=lambda mid: None, shuffle_options=True):
    """
    questions: [{stem, stem_xml:[...], options:[{text,xml:[...],correct?}], correct_index}]
    get_media(mid)->bytes : rasm baytlarini olish uchun (DBMEDIA tokenlari)
    return: (test_bytes, key_bytes)
    """
    test = Document(); _set_default_font(test)
    key = Document(); _set_default_font(key)

    tp = test.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title); r.bold = True; r.font.size = Pt(15)
    test.add_paragraph()
    kp = key.add_paragraph(); kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kp.add_run(title + " — JAVOBLAR"); r.bold = True; r.font.size = Pt(15)
    key.add_paragraph()

    key_line = []
    for i, q in enumerate(questions):
        opts = list(enumerate(q["options"]))  # (orig_index, option)
        if shuffle_options:
            random.shuffle(opts)

        p = test.add_paragraph()
        rr = p.add_run("#%d. " % (i + 1)); rr.bold = True
        _append_fragments(test, p, q.get("stem_xml", []), get_media)

        correct_letter = "?"
        for j, (orig_idx, opt) in enumerate(opts):
            op = test.add_paragraph()
            lr = op.add_run("%s) " % LETTERS[j]); lr.bold = True
            _append_fragments(test, op, opt.get("xml", []), get_media)
            if orig_idx == q["correct_index"]:
                correct_letter = LETTERS[j]
        test.add_paragraph()

        key_line.append("%d-%s" % (i + 1, correct_letter))
        if len(key_line) == 10 or i == len(questions) - 1:
            key.add_paragraph("     ".join(key_line))
            key_line = []

    tb, kb = io.BytesIO(), io.BytesIO()
    test.save(tb); key.save(kb)
    return tb.getvalue(), kb.getvalue()
